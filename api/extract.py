import io
import json
import base64
from http.server import BaseHTTPRequestHandler


def extract_text(filename, data_bytes):
    """Extrae texto plano de un PDF, DOCX o TXT. Devuelve string (puede ser vacío)."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data_bytes))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    if name.endswith(".docx"):
        import docx
        doc = docx.Document(io.BytesIO(data_bytes))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    # .txt o desconocido: intentar decodificar como texto
    return data_bytes.decode("utf-8", errors="replace")


def process(body):
    """Recibe {filename, data(base64)} y devuelve (status, {text, chars} | {error|warning})."""
    filename = body.get("filename", "")
    b64 = body.get("data", "")
    if not b64:
        return 400, {"error": "No se recibió ningún archivo."}

    name = (filename or "").lower()
    if not name.endswith((".pdf", ".docx", ".txt")):
        return 400, {"error": "Formato no soportado. Subí un PDF, DOCX o TXT."}

    try:
        # El front puede mandar un data URL (data:...;base64,XXXX) o base64 puro
        if b64.strip().startswith("data:") and "," in b64:
            b64 = b64.split(",", 1)[1]
        raw = base64.b64decode(b64)
    except Exception as e:
        return 400, {"error": f"Archivo inválido: {str(e)}"}

    try:
        text = extract_text(filename, raw).strip()
    except Exception as e:
        return 500, {"error": f"No se pudo leer el archivo: {str(e)}"}

    if not text:
        return 200, {
            "text": "",
            "warning": "No se pudo extraer texto. Si es un CV escaneado como imagen, "
                       "necesita OCR o texto seleccionable."
        }

    return 200, {"text": text, "chars": len(text)}


class handler(BaseHTTPRequestHandler):

    def log_message(self, format, *args):
        pass

    def send_json(self, status, data):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_POST(self):
        try:
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length).decode("utf-8"))
            status, result = process(body)
            self.send_json(status, result)
        except Exception as e:
            self.send_json(500, {"error": str(e)})
